from flask import Flask, render_template, request
import datetime
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

def data_processing():
    global pgm_title, venue, address, date, time, resource_person, collab_agency, level, no_teach, no_stud, short_write_up, skills, link, signatures, name_doc_crtr, phn_doc_crtr, dept
    pgm_title = 'Program Title'
    venue = 'Auditoium'
    date = str(datetime.date(2021, 12, 31))
    time = str(datetime.datetime.fromtimestamp(1630425600).strftime('%Y-%m-%d %H:%M:%S'))
    resource_person = "John Doe"
    collab_agency = "XYZ"
    level = "National"
    no_teach = str(10)
    no_stud = str(100)
    short_write_up = "This is a short write-up"
    skills = "Skills"
    link = "www.google.com"
    signatures = "Signatures"
    name_doc_crtr = "John Doe"
    phn_doc_crtr = "1234567890"
    dept = "Computer Science"

def image_processing():
    global brochure, cert, attnd_std, attnd_teach, pic1, pic2, pic3
    brochure = 'test.png'
    cert = 'test.png'
    attnd_std = 'test.png'
    attnd_teach = 'test.png'
    pic1 = 'test.png'
    pic2 = 'test.png'
    pic3 = 'test.png'

def replace_text_placeholder(doc, placeholder, replacement, bold=False, font_size=None):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
            # Apply formatting to the new text
            for run in paragraph.runs:
                # Apply formatting
                run.font.name = 'Times New Roman'
                if bold:
                    run.font.bold = True
                if font_size is not None:
                    run.font.size = Pt(font_size)

""" def replace_placeholder(doc, placeholder, replacement, bold=False, font_size=None):
    for table in doc.tables:
        # Iterate over rows
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, replacement) """

def replace_table_placeholder(doc, placeholder, replacement, bold=False, font_size=None):
    for table in doc.tables:
        # Iterate over rows
        for row in table.rows:
            for cell in row.cells:
                # Check if placeholder exists in cell text
                if placeholder in cell.text:
                    # Replace placeholder in cell text
                    cell.text = cell.text.replace(placeholder, replacement)
                    # Apply formatting to the replaced text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            # Apply formatting
                            run.font.name = 'Times New Roman'
                            if bold:
                                run.font.bold = True
                            if font_size is not None:
                                run.font.size = Pt(font_size)


def replace_image_placeholder(doc, placeholder, image_path, width):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = ''
                    # Add space above the image
                    cell.paragraphs[0].add_run().add_break()
                    paragraph = cell.paragraphs[0].add_run()
                    image_path = 'test.png'  # Replace 'test.png' with the path to your image
                    width = Cm(width)  # Adjust the width as needed to fit in the cell
                    paragraph.add_picture(image_path, width=width)
                    # Add space below the image
                    cell.paragraphs[0].add_run().add_break()
                    # Center align the image in the cell
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def replace_placeholder_in_footer(doc, placeholder, replacement):
    for section in doc.sections:
        footer = section.footer
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, replacement)

def generate_report():
    # Load Word template
    doc = Document('Report Template.docx')
    
    replace_text_placeholder(doc,'placeholder#pgm-title',pgm_title, True,24)

    # Placeholder replacements for text
    replace_table_placeholder(doc, 'placeholder#venue', venue)
    replace_table_placeholder(doc, 'placeholder#date', date)
    replace_table_placeholder(doc, 'placeholder#time', time)
    replace_table_placeholder(doc, 'placeholder#resource-person', resource_person)
    replace_table_placeholder(doc, 'placeholder#collab-agency', collab_agency)
    replace_table_placeholder(doc, 'placeholder#level', level)
    replace_table_placeholder(doc, 'placeholder#no-teach', no_teach)
    replace_table_placeholder(doc, 'placeholder#no-stud', no_stud)
    replace_table_placeholder(doc, 'placeholder#short-write-up', short_write_up)
    replace_table_placeholder(doc, 'placeholder#skills', skills)
    replace_table_placeholder(doc, 'placeholder#link', link)
    replace_table_placeholder(doc, 'placeholder#signatures', signatures)
    replace_table_placeholder(doc, 'placeholder#name-doc-crtr', name_doc_crtr)
    replace_table_placeholder(doc, 'placeholder#phn-doc-crtr', phn_doc_crtr)

    # Placeholder replacements for images
    replace_image_placeholder(doc, 'images#brochure', brochure,10)
    replace_image_placeholder(doc, 'images#cert', cert,15)
    replace_image_placeholder(doc, 'images#attnd-stud', attnd_std,10)
    replace_image_placeholder(doc, 'images#attnd-teach', attnd_teach,10)
    replace_image_placeholder(doc, 'images#pic1', pic1,4)
    replace_image_placeholder(doc, 'images#pic2', pic2,4)
    replace_image_placeholder(doc, 'images#pic3', pic3,4)

    replace_placeholder_in_footer(doc, 'Placeholder#Dept', dept)

    # Save modified Word document
    doc.save(pgm_title+'_Report.docx')
    convert('output.docx', pgm_title+'_Report.pdf')

    return 'PDF generated successfully!'

data_processing()
image_processing()
generate_report()


