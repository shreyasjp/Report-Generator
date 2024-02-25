from flask import Flask, render_template, request
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def process_form():
    # Load Word template
    doc = Document('Report Template.docx')

    # Find and replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        if 'placeholder#pgm-title' in paragraph.text:
            paragraph.text = paragraph.text.replace('placeholder#pgm-title', 'tutu')
            # Apply formatting to the new text
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(24)
                run.font.bold = True
        # Add more placeholders as needed

    # Find and replace placeholders in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if 'images#pic1' in cell.text:
                    cell.text = ''
                    # Add space above the image
                    cell.paragraphs[0].add_run().add_break()
                    paragraph = cell.paragraphs[0].add_run()
                    image_path = 'test.png'  # Replace 'test.png' with the path to your image
                    width = Cm(4)  # Adjust the width as needed to fit in the cell
                    paragraph.add_picture(image_path, width=width)
                    # Add space below the image
                    cell.paragraphs[0].add_run().add_break()
                    # Center align the image in the cell
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Save modified Word document
    doc.save('output.docx')

    return 'PDF generated successfully!'


process_form()
