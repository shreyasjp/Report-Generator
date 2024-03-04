import tkinter as tk
from tkinter import filedialog, messagebox
import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

def generate_report():
    global pgm_title, venue, date, time, resource_person, collab_agency, level, no_teach, no_stud, short_write_up, skills, link, signatures, name_doc_crtr, phn_doc_crtr, dept
    pgm_title = program_title_entry.get()
    venue = venue_entry.get()
    date = date_entry.get()
    time = time_entry.get()
    resource_person = resource_person_entry.get()
    collab_agency = collab_agency_entry.get()
    level = level_entry.get()
    no_teach = no_teach_entry.get()
    no_stud = no_stud_entry.get()
    short_write_up = short_write_up_entry.get()
    skills = skills_entry.get()
    link = link_entry.get()
    signatures = signatures_entry.get()
    name_doc_crtr = name_doc_crtr_entry.get()
    phn_doc_crtr = phn_doc_crtr_entry.get()
    dept = dept_entry.get()
    
    # Validate required fields
    if not all([pgm_title, venue, date, time, resource_person, collab_agency, level, no_teach, no_stud, short_write_up, skills, link, signatures, name_doc_crtr, phn_doc_crtr, dept]):
        messagebox.showerror("Error", "Please fill in all required fields")
        return

    # Load Word template
    doc = Document('Report Template.docx')

    replace_text_placeholder(doc, 'placeholder#pgm-title', pgm_title, True, 24)
    replace_table_placeholder(doc, 'placeholder#venue', venue)
    replace_table_placeholder(doc, 'placeholder#date', date)
    replace_table_placeholder(doc, 'placeholder#time', time)
    replace_table_placeholder(doc, 'placeholder#resource-person', resource_person)
    replace_table_placeholder(doc, 'placeholder#collab-agency', collab_agency)
    replace_table_placeholder(doc, 'placeholder#level', level)
    replace_table_placeholder(doc, 'placeholder#no-teach', no_teach)
    replace_table_placeholder(doc, 'placeholder#no-stud', no_stud)
    replace_table_placeholder(doc, 'placeholder#short-write-up', short_write_up)
    replace_table_placeholder(doc, 'placeholder#skills', skills)
    replace_table_placeholder(doc, 'placeholder#link', link)
    replace_table_placeholder(doc, 'placeholder#signatures', signatures)
    replace_table_placeholder(doc, 'placeholder#name-doc-crtr', name_doc_crtr)
    replace_table_placeholder(doc, 'placeholder#phn-doc-crtr', phn_doc_crtr)

    replace_placeholder_in_footer(doc, 'Placeholder#Dept', dept)

    # Placeholder replacements for images
    for placeholder, image_path in images.items():
        if image_path:
            image_replace(doc, placeholder, image_path, 10)

    # Save modified Word document
    report_name = pgm_title + '_Report.docx'
    doc.save(report_name)
    convert(report_name, pgm_title + '_Report.pdf')

    messagebox.showinfo("Success", "PDF generated successfully!")

def image_replace(doc, placeholder, image_path, width):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.clear()
                        paragraph.add_run().add_break()
                        run = paragraph.add_run()
                        run.add_picture(image_path, width=Cm(width))
                        paragraph.add_run().add_break()
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def replace_text_placeholder(doc, placeholder, replacement, bold=False, font_size=None):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                if bold:
                    run.font.bold = True
                if font_size is not None:
                    run.font.size = Pt(font_size)

def replace_table_placeholder(doc, placeholder, replacement, bold=False, font_size=None):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, replacement)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            if bold:
                                run.font.bold = True
                            if font_size is not None:
                                run.font.size = Pt(font_size)

def replace_placeholder_in_footer(doc, placeholder, replacement):
    for section in doc.sections:
        footer = section.footer
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, replacement)

def select_image(placeholder):
    filename = filedialog.askopenfilename(initialdir="/", title="Select Image", filetypes=(("PNG files", "*.png"), ("JPEG files", "*.jpg"), ("All files", "*.*")))
    if filename:
        images[placeholder] = filename

# Create GUI window
images = {
    "images#brochure": "",
    "images#cert": "",
    "images#attnd-stud": "",
    "images#attnd-teach": "",
    "images#pic1": "",
    "images#pic2": "",
    "images#pic3": ""
}
root = tk.Tk()
root.title("Report Generation")

# Create input fields
fields = [
    ("Program Title:", "program_title_entry"),
    ("Venue:", "venue_entry"),
    ("Date:", "date_entry"),
    ("Time:", "time_entry"),
    ("Resource Person:", "resource_person_entry"),
    ("Collaborating Agency:", "collab_agency_entry"),
    ("Level:", "level_entry"),
    ("No. of Teachers:", "no_teach_entry"),
    ("No. of Students:", "no_stud_entry"),
    ("Short Write-Up:", "short_write_up_entry"),
    ("Skills:", "skills_entry"),
    ("Link:", "link_entry"),
    ("Signatures:", "signatures_entry"),
    ("Name of Document Creator:", "name_doc_crtr_entry"),
    ("Phone of Document Creator:", "phn_doc_crtr_entry"),
    ("Department:", "dept_entry")
]

for i, (label, entry_name) in enumerate(fields):
    lbl = tk.Label(root, text=label)
    lbl.grid(row=i, column=0, sticky='e')
    entry = tk.Entry(root)
    entry.grid(row=i, column=1, padx=5, pady=5)
    exec(f"{entry_name} = entry")

# Buttons to select images
buttons = [
    ("Brochure", "images#brochure"),
    ("Certificate", "images#cert"),
    ("Attendance (Student)", "images#attnd-stud"),
    ("Attendance (Teacher)", "images#attnd-teach"),
    ("Pic1", "images#pic1"),
    ("Pic2", "images#pic2"),
    ("Pic3", "images#pic3")
]

for i, (button_label, placeholder) in enumerate(buttons):
    btn = tk.Button(root, text=f"Select {button_label}", command=lambda ph=placeholder: select_image(ph))
    btn.grid(row=i, column=2, padx=5, pady=5)

# Button to generate report
generate_button = tk.Button(root, text="Generate Report", command=generate_report)
generate_button.grid(row=len(fields), columnspan=3, pady=10)

root.mainloop()
