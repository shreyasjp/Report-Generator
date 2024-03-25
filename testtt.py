import tkinter as tk
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageTk

doc = Document('Report Template.docx')

def replace_paragraph(placeholder, replacement, font='Times New Roman', size = 12, bold = False):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
            for run in paragraph.runs:
                run.font.name = font
                run.font.size = Pt(size)
                run.font.bold = bold

# replace_paragraph('placeholder#pgm-title', 'Greeshma Shajan', size = 24, bold = True)

def replace_table_data(placeholder,replacement,font='Times New Roman',size=12):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, replacement)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = font
                            run.font.size = Pt(size)

# replace_table_data('placeholder#venue', 'Kochi')

def replace_table_image(placeholder,path,width):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.clear()
                        paragraph.add_run().add_break()
                        run = paragraph.add_run()
                        run.add_picture(path, width=Cm(width))
                        paragraph.add_run().add_break()
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# replace_table_image('images#pic2', 'Ms. Nkita Pinheiro.png', 10)
                        
def process_form():
    global textbox_1, input_1

    a = textbox_1.get("1.0", "end-1c")
    b = input_1.get()

    replace_table_data('placeholder#short-write-up', a)
    replace_paragraph('placeholder#pgm-title', b, size = 24, bold = True)

    doc.save('Greeshma.docx')

root = tk.Tk()
root.title("Report")

ico = Image.open('WhatsApp Image 2024-02-20 at 16.25.55_77aad384.jpg')
photo = ImageTk.PhotoImage(ico)
root.wm_iconphoto(False, photo)

root.geometry("600x400")

label_1 = tk.Label(root, text="Enter Title:")
label_1.grid(row=0, column=0)

textbox_1 = tk.Text(root, height=5, width=20)
textbox_1.grid(row=0, column=1)

label_2 = tk.Label(root, text='Enter Value 1')
label_2.grid(row=1, column=0)

input_1 = tk.Entry(root)
input_1.grid(row=1, column=1)

button = tk.Button(root, text="Submit", command=process_form)
button.grid(row=2, column=1)

root.mainloop()