from flask import Flask, render_template, request, redirect, url_for
import datetime
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

app = Flask(__name__)

# Default values for form fields
default_values = {
    'pgm_title': 'NA',
    'venue': 'NA',
    'date': 'NA',
    'time': 'NA',
    'resource_person': 'NA',
    'collab_agency': 'NA',
    'level': 'NA',
    'no_teach': 'NA',
    'no_stud': 'NA',
    'short_write_up': 'NA',
    'skills': 'NA',
    'link': 'NA',
    'signatures': 'NA',
    'name_doc_crtr': 'NA',
    'phn_doc_crtr': 'NA',
    'dept': 'NA',
}

def generate_report(form_data):
    # Load Word template
    doc = Document('Report Template.docx')
    
    # Placeholder replacements for text
    replace_text_placeholder(doc, 'placeholder#pgm-title', form_data['pgm_title'], True, 24)
    replace_table_placeholder(doc, 'placeholder#venue', form_data['venue'])
    replace_table_placeholder(doc, 'placeholder#date', form_data['date'])
    replace_table_placeholder(doc, 'placeholder#time', form_data['time'])
    replace_table_placeholder(doc, 'placeholder#resource-person', form_data['resource_person'])
    replace_table_placeholder(doc, 'placeholder#collab-agency', form_data['collab_agency'])
    replace_table_placeholder(doc, 'placeholder#level', form_data['level'])
    replace_table_placeholder(doc, 'placeholder#no-teach', form_data['no_teach'])
    replace_table_placeholder(doc, 'placeholder#no-stud', form_data['no_stud'])
    replace_table_placeholder(doc, 'placeholder#short-write-up', form_data['short_write_up'])
    replace_table_placeholder(doc, 'placeholder#skills', form_data['skills'])
    replace_table_placeholder(doc, 'placeholder#link', form_data['link'])
    replace_table_placeholder(doc, 'placeholder#signatures', form_data['signatures'])
    replace_table_placeholder(doc, 'placeholder#name-doc-crtr', form_data['name_doc_crtr'])
    replace_table_placeholder(doc, 'placeholder#phn-doc-crtr', form_data['phn_doc_crtr'])

    # Placeholder replacements for images
    """ replace_image_placeholder(doc, 'images#brochure', form_data['brochure'], 10)
    replace_image_placeholder(doc, 'images#cert', form_data['cert'], 15)
    replace_image_placeholder(doc, 'images#attnd-stud', form_data['attnd_std'], 10)
    replace_image_placeholder(doc, 'images#attnd-teach', form_data['attnd_teach'], 10)
    replace_image_placeholder(doc, 'images#pic1', form_data['pic1'], 4)
    replace_image_placeholder(doc, 'images#pic2', form_data['pic2'], 4)
    replace_image_placeholder(doc, 'images#pic3', form_data['pic3'], 4) """

    replace_placeholder_in_footer(doc, 'Placeholder#Dept', form_data['dept'])

    # Save modified Word document
    file_name = form_data['pgm_title'] + '_Report.docx'
    doc.save(file_name)
    convert(file_name, file_name.replace('.docx', '.pdf'))

@app.route('/')
def index():
    return render_template('form.html', default_values=default_values)

@app.route('/submit', methods=['POST'])
def submit():
    # Extract form data
    form_data = {key: request.form[key] for key in default_values.keys()}
    generate_report(form_data)
    return redirect(url_for('index'))

def replace_text_placeholder(doc, placeholder, replacement, bold=False, font_size=None):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
            # Apply formatting to the new text
            for run in paragraph.runs:
                # Apply formatting
                run.font.name = 'Times New Roman'
                if bold:
                    run.font.bold = True
                if font_size is not None:
                    run.font.size = Pt(font_size)

def replace_table_placeholder(doc, placeholder, replacement, bold=False, font_size=None):
    for table in doc.tables:
        # Iterate over rows
        for row in table.rows:
            for cell in row.cells:
                # Check if placeholder exists in cell text
                if placeholder in cell.text:
                    # Replace placeholder in cell text
                    cell.text = cell.text.replace(placeholder, replacement)
                    # Apply formatting to the replaced text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            # Apply formatting
                            run.font.name = 'Times New Roman'
                            if bold:
                                run.font.bold = True
                            if font_size is not None:
                                run.font.size = Pt(font_size)

def replace_image_placeholder(doc, placeholder, image_path, width):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = ''
                    # Add space above the image
                    cell.paragraphs[0].add_run().add_break()
                    paragraph = cell.paragraphs[0].add_run()
                    width = Cm(width)  # Adjust the width as needed to fit in the cell
                    paragraph.add_picture(image_path, width=width)
                    # Add space below the image
                    cell.paragraphs[0].add_run().add_break()
                    # Center align the image in the cell
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def replace_placeholder_in_footer(doc, placeholder, replacement):
    for section in doc.sections:
        footer = section.footer
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, replacement)

if __name__ == '__main__':
    app.run(debug=True)
